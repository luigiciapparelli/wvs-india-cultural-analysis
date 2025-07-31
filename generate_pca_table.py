import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Replace with your actual CSV path if different
df = pd.read_csv("Data/pca_variance_alpha.csv")

# Take top 5 rows (or change this number as needed)
df = df.head(5)

# Round numeric values
df["pve"] = df["pve"].round(5)
df["alpha"] = df["alpha"].round(5)

# Create Word document
doc = Document()
doc.add_heading("PCA Combinations – Top Results", level=1)

table = doc.add_table(rows=1, cols=len(df.columns))
table.style = 'Table Grid'

# Set column headers
hdr_cells = table.rows[0].cells
for i, col_name in enumerate(df.columns):
    hdr_cells[i].text = col_name
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.font.bold = True
    run.font.size = Pt(10)

# Fill in data
for _, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)
        run = row_cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(10)

# Save to file
doc.save("pca_combinations_table_styled.docx")
print("Styled PCA table saved as pca_combinations_table_styled.docx")
